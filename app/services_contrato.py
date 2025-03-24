import os
import re
from docx import Document
from docx.shared import Pt

# Regex para identificar os placeholders, ignorando espaços extras
PLACEHOLDER_PATTERN = re.compile(r'{{\s*(.*?)\s*}}')

# Lista de trechos especiais que devem permanecer em negrito (se já estiverem formatados no template)
SPECIAL_BOLD = ["CLÁUSULA 2ª:", "PIX OU LINK DE PARCELAMENTO"]

def format_phone(phone_str):
    """
    Remove caracteres não numéricos e formata o número de telefone.
    Se tiver 11 dígitos: formata como "DD NNNNN-NNNN".
    Se tiver 10 dígitos: formata como "DD NNNN-NNNN".
    Caso contrário, retorna o número sem formatação.
    """
    digits = re.sub(r'\D', '', phone_str)
    if len(digits) == 11:
        return f"{digits[:2]} {digits[2:7]}-{digits[7:]}"
    elif len(digits) == 10:
        return f"{digits[:2]} {digits[2:6]}-{digits[6:]}"
    else:
        return phone_str

def merge_paragraph_runs(paragraph):
    """
    Mescla todos os runs do parágrafo em um único run,
    garantindo que os placeholders fiquem íntegros.
    """
    full_text = paragraph.text
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)
    new_run = paragraph.add_run(full_text)
    return new_run

def process_run(run, dados_proc):
    original_text = run.text
    if "{{" not in original_text:
        return [(original_text, run.bold, run.underline, run.style, run.font.name, run.font.size)]
    
    tokens = re.split(PLACEHOLDER_PATTERN, original_text)
    output = []
    is_placeholder = False
    for token in tokens:
        if is_placeholder:
            key = token.strip().lower()
            val = dados_proc.get(key, "")
            # Se o placeholder for ddd_telefone, force a fonte Times New Roman
            if key == "ddd_telefone":
                output.append((val, True, run.underline, run.style, "Times New Roman", Pt(12)))
            else:
                output.append((val, True, run.underline, run.style, run.font.name, run.font.size))
        else:
            output.append((token, run.bold, run.underline, run.style, run.font.name, run.font.size))
        is_placeholder = not is_placeholder
    return output

def process_paragraph(paragraph, dados_proc):
    """
    Processa um parágrafo do contrato:
      - Se o parágrafo contiver "GUARULHUS" (ignore case) e "2025", força todo o parágrafo a ficar em bold.
      - Caso contrário, verifica se o placeholder está fragmentado e, se necessário, mescla os runs.
      - Em seguida, processa run a run, substituindo os placeholders e preservando a formatação original.
      - Tokens literais que correspondem a SPECIAL_BOLD serão forçados a ficar em bold.
      - Após reconstruir o parágrafo, aplica um fallback para substituir quaisquer placeholders remanescentes.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" in full_text and not any("{{" in run.text and "}}" in run.text for run in paragraph.runs):
        merge_paragraph_runs(paragraph)
        full_text = paragraph.text

    if "GUARULHUS" in full_text.upper() and "2025" in full_text:
        text_subst = re.sub(PLACEHOLDER_PATTERN,
                            lambda m: dados_proc.get(m.group(1).strip().lower(), ""),
                            full_text)
        paragraph.text = ""
        run = paragraph.add_run(text_subst)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        return

    if "{{" not in full_text:
        return

    all_tokens = []
    for run in paragraph.runs:
        tokens = process_run(run, dados_proc)
        all_tokens.extend(tokens)
    
    processed_tokens = []
    for text, bold, underline, style, font_name, font_size in all_tokens:
        if text.strip().upper() in (s.upper() for s in SPECIAL_BOLD):
            processed_tokens.append((text, True, underline, style, font_name, font_size))
        else:
            processed_tokens.append((text, bold, underline, style, font_name, font_size))
    
    paragraph.text = ""
    for token_text, is_bold, underline, style, font_name, font_size in processed_tokens:
        new_run = paragraph.add_run(token_text)
        new_run.bold = is_bold
        new_run.underline = underline
        if font_name:
            new_run.font.name = font_name
        if font_size:
            new_run.font.size = font_size

    for run in paragraph.runs:
        if "{{" in run.text:
            run.text = re.sub(PLACEHOLDER_PATTERN,
                              lambda m: dados_proc.get(m.group(1).strip().lower(), ""),
                              run.text)
    
    # Fallback específico para o placeholder "ddd_telefone"
    for run in paragraph.runs:
        if "{{ddd_telefone}}" in run.text:
            run.text = run.text.replace("{{ddd_telefone}}", dados_proc.get("ddd_telefone", ""))
            run.bold = True
            run.font.name = "Times New Roman"

def preencher_contrato(dados, base_dir):
    """
    Gera o contrato com base no modelo (arquivo2.docx).
      - Converte os valores dos campos para maiúsculas (exceto "mes", que é capitalizado).
      - As chaves dos dados são convertidas para minúsculas.
      - Apenas os valores que substituem os placeholders ficam em bold, preservando a formatação original do documento.
      - Se o parágrafo contiver "GUARULHUS" e "2025", a linha inteira fica em bold.
      - SPECIAL_BOLD permanecem conforme no template.
    """
    modelo_path = os.path.join(base_dir, "..", "templates", "arquivo2.docx")
    if not os.path.exists(modelo_path):
        raise FileNotFoundError(f"Modelo não encontrado: {modelo_path}")
    
    doc = Document(modelo_path)
    
    # Prepara os dados: converte os valores para maiúsculas (exceto "mes" que é capitalizado)
    dados_proc = {}
    for chave, valor in dados.items():
        if valor:
            if chave.lower() == "mes":
                dados_proc[chave.lower()] = valor.strip().capitalize()
            elif chave.lower() == "ddd_telefone":
                dados_proc[chave.lower()] = format_phone(valor.strip())
            else:
                dados_proc[chave.lower()] = valor.strip().upper()
        else:
            dados_proc[chave.lower()] = ""
    
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, dados_proc)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph, dados_proc)
    
    output_dir = os.path.join(base_dir, "..", "output")
    os.makedirs(output_dir, exist_ok=True)
    nome_cliente = dados_proc.get("nome", "SEM_NOME")
    doc_path = os.path.join(output_dir, f"CONTRATO_{nome_cliente}.docx".replace(" ", "_"))
    doc.save(doc_path)
    return doc_path
