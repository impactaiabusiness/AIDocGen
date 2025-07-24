import os
import re
from docx import Document
from docx.shared import Pt

PLACEHOLDER_PATTERN = re.compile(r'{{(.*?)}}')

def substituir_placeholders_procuracao(paragraph, dados):
    """
    Substitui os placeholders no parágrafo.
    - Os valores inseridos (aqueles que substituem os placeholders) serão formatados em negrito;
    - Se o parágrafo for o do outorgante (ou seja, começar com "OUTORGANTE:"), todo o parágrafo ficará em negrito.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    outorgante_flag = full_text.strip().startswith("OUTORGANTE:")

    if "{{" not in full_text:
        if outorgante_flag:
            for run in paragraph.runs:
                run.bold = True
        return

    parts = re.split(PLACEHOLDER_PATTERN, full_text)
    new_tokens = []
    is_placeholder = False
    for part in parts:
        if is_placeholder:
            valor = dados.get(part.strip(), "")
            new_tokens.append((valor, True))
        else:
            new_tokens.append((part, False))
        is_placeholder = not is_placeholder

    paragraph.text = ""
    for token_text, is_bold in new_tokens:
        run = paragraph.add_run(token_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = True if outorgante_flag else is_bold

def preencher_procuracao_jacquelline(dados, base_dir):
    """
    Gera a procuração da Dra. Jaquelline com base no template arquivo1.docx.
    """
    modelo_path = os.path.join(base_dir, "..", "templates", "arquivo1.docx")
    return _gerar_procuracao(dados, modelo_path, base_dir)

def preencher_procuracao_jose(dados, base_dir):
    """
    Gera a procuração do Dr. José com base no template arquivo2.docx.
    """
    modelo_path = os.path.join(base_dir, "..", "templates", "arquivo2.docx")
    return _gerar_procuracao(dados, modelo_path, base_dir)

def _gerar_procuracao(dados, modelo_path, base_dir):
    """
    Função interna que faz o processamento comum para gerar procurações.
    """
    if not os.path.exists(modelo_path):
        raise FileNotFoundError(f"Modelo não encontrado: {modelo_path}")
    
    doc = Document(modelo_path)
    
    dados_processados = {chave: valor.strip().upper() if valor else "" for chave, valor in dados.items()}
    
    for para in doc.paragraphs:
        substituir_placeholders_procuracao(para, dados_processados)
    
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for para in celula.paragraphs:
                    substituir_placeholders_procuracao(para, dados_processados)
    
    output_dir = os.path.join(base_dir, "..", "output")
    os.makedirs(output_dir, exist_ok=True)
    nome_cliente = dados_processados.get("nome", "SEM_NOME")
    doc_path = os.path.join(output_dir, f"PROCURACAO_{nome_cliente}.docx".replace(" ", "_"))
    doc.save(doc_path)
    return doc_path
